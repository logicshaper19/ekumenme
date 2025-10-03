"""
Knowledge Base Contribution Workflow Service
Implements the vetted, time-bound knowledge contribution system
"""

import asyncio
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple
from enum import Enum
from pathlib import Path
import uuid
import json

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_, or_, func
from sqlalchemy.orm import selectinload

from app.core.database import get_async_db
from app.models.organization import Organization
from app.models.user import User
from app.models.knowledge_base import KnowledgeBaseDocument, DocumentStatus, DocumentType
from app.services.optimized_llm_service import OptimizedLLMService
from app.agents.knowledge_base_compliance_agent import KnowledgeBaseComplianceAgent
try:
    from app.services.notification_service import NotificationService
except ImportError:
    # Mock notification service for testing
    class NotificationService:
        async def send_notification(self, **kwargs):
            logger.info(f"Mock notification: {kwargs}")
    NotificationService = NotificationService
from app.core.config import settings

logger = logging.getLogger(__name__)


class SubmissionStatus(str, Enum):
    """Status of knowledge base submissions"""
    SUBMITTED = "submitted"
    UNDER_REVIEW = "under_review"
    APPROVED = "approved"
    REJECTED = "rejected"
    EXPIRED = "expired"
    RENEWAL_PENDING = "renewal_pending"


class ContentQualityScore:
    """Content quality assessment results"""
    
    def __init__(self, score: float, issues: List[str], recommendations: List[str]):
        self.score = score  # 0-100
        self.issues = issues
        self.recommendations = recommendations
    
    def is_acceptable(self, threshold: float = 70.0) -> bool:
        return self.score >= threshold


class KnowledgeBaseWorkflowService:
    """Service for managing the knowledge base contribution workflow"""
    
    def __init__(self):
        self.llm_service = OptimizedLLMService()
        self.notification_service = NotificationService()
        self.compliance_agent = KnowledgeBaseComplianceAgent()
        
        # Configuration
        self.default_expiration_months = 18  # 18 months default
        self.reminder_days = [90, 60, 30]  # Days before expiration
        self.quality_threshold = 70.0  # Minimum quality score for auto-approval (legacy)
        
    async def submit_document(
        self,
        organization_id: str,
        uploaded_by: str,
        file_path: str,
        filename: str,
        file_type: str,
        document_type: DocumentType,
        description: Optional[str] = None,
        tags: Optional[List[str]] = None,
        expiration_months: Optional[int] = None,
        visibility: str = "internal",
        shared_with_organizations: Optional[List[str]] = None,
        shared_with_users: Optional[List[str]] = None,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """
        Step 1: Document Submission Portal
        Process initial document submission
        """
        try:
            logger.info(f"Processing document submission: {filename} from org {organization_id}")
            
            # Create document record
            document_id = uuid.uuid4()
            expiration_date = datetime.utcnow() + timedelta(
                days=(expiration_months or self.default_expiration_months) * 30
            )
            
            document = KnowledgeBaseDocument(
                id=document_id,
                organization_id=organization_id,
                uploaded_by=uploaded_by,
                filename=filename,
                file_path=file_path,
                file_type=file_type,
                file_size_bytes=Path(file_path).stat().st_size if Path(file_path).exists() else 0,
                processing_status=DocumentStatus.PENDING,
                document_type=document_type,
                tags=tags or [],
                description=description,
                visibility=visibility,  # Use provided visibility setting
                shared_with_organizations=shared_with_organizations or [],
                shared_with_users=shared_with_users or [],
                is_ekumen_provided=False,
                submission_status="submitted",
                expiration_date=expiration_date,
                auto_renewal=False,
                quality_score=0.0,
                quality_issues=[],
                quality_recommendations=[],
                version=1,
                workflow_metadata={
                    "submitted_at": datetime.utcnow().isoformat(),
                    "expiration_months": expiration_months or self.default_expiration_months
                }
            )
            
            db.add(document)
            await db.commit()
            
            # Step 2: Regulatory Compliance Validation
            compliance_result = await self._validate_regulatory_compliance(file_path, document_type, str(document_id))
            
            # Update document with compliance results
            document.processing_status = DocumentStatus.PROCESSING
            document.submission_status = "under_review"
            
            # Store compliance results in workflow metadata
            document.workflow_metadata.update({
                "compliance_validation_completed": datetime.utcnow().isoformat(),
                "compliance_decision": compliance_result.get("decision", "UNCERTAIN"),
                "compliance_confidence": compliance_result.get("confidence", 0.0),
                "violations": compliance_result.get("violations", []),
                "warnings": compliance_result.get("warnings", []),
                "recommendations": compliance_result.get("recommendations", []),
                "tool_results": compliance_result.get("tool_results", {})
            })
            
            await db.commit()
            
            # Make approval decision based on compliance
            decision = compliance_result.get("decision", "uncertain")
            if decision == "auto_approve":
                await self._auto_approve_document(str(document_id), db)
                submission_status = SubmissionStatus.APPROVED
            elif decision == "flag_for_review":
                submission_status = SubmissionStatus.UNDER_REVIEW
                # Notify reviewers
                await self._notify_reviewers(str(document_id), compliance_result, db)
            else:  # uncertain
                submission_status = SubmissionStatus.UNDER_REVIEW
                await self._notify_reviewers(str(document_id), compliance_result, db)
            
            return {
                "document_id": str(document_id),
                "status": submission_status,
                "compliance_decision": compliance_result.get("decision", "UNCERTAIN"),
                "compliance_confidence": compliance_result.get("confidence", 0.0),
                "violations": compliance_result.get("violations", []),
                "warnings": compliance_result.get("warnings", []),
                "recommendations": compliance_result.get("recommendations", []),
                "expiration_date": expiration_date.isoformat(),
                "message": self._get_submission_message(submission_status, compliance_result)
            }
            
        except Exception as e:
            logger.error(f"Error submitting document: {e}")
            raise
    
    async def _validate_regulatory_compliance(
        self,
        file_path: str,
        document_type: DocumentType,
        document_id: str
    ) -> Dict[str, Any]:
        """
        Validate document for regulatory compliance using the compliance agent
        
        Args:
            file_path: Path to the document file
            document_type: Type of document
            document_id: Document identifier
            
        Returns:
            Compliance validation result
        """
        try:
            logger.info(f"Starting regulatory compliance validation for document {document_id}")
            
            # Extract text content from file
            content = await self._extract_text_content(file_path)
            
            # Use compliance agent to validate
            compliance_result = await self.compliance_agent.validate_document(
                document_content=content,
                document_type=document_type.value,
                document_id=document_id
            )
            
            logger.info(f"Compliance validation completed: {compliance_result.get('decision', 'UNKNOWN')}")
            
            return compliance_result
            
        except Exception as e:
            logger.error(f"Error in regulatory compliance validation: {e}")
            return {
                "decision": "UNCERTAIN",
                "confidence": 0.0,
                "reason": f"Compliance validation failed: {str(e)}",
                "violations": [f"Validation error: {str(e)}"],
                "warnings": [],
                "recommendations": ["Manual review required due to validation error"],
                "tool_results": {},
                "timestamp": datetime.utcnow().isoformat()
            }
    
    def _get_submission_message(self, submission_status: str, compliance_result: Dict[str, Any]) -> str:
        """Get appropriate submission message based on status and compliance result"""
        decision = compliance_result.get("decision", "uncertain")
        violations = compliance_result.get("violations", [])
        warnings = compliance_result.get("warnings", [])
        
        if decision == "auto_approve":
            return "Document auto-approved - fully compliant with French regulations"
        elif decision == "flag_for_review":
            if violations:
                return f"Document flagged for review - {len(violations)} regulatory violations found"
            elif warnings:
                return f"Document flagged for review - {len(warnings)} compliance warnings"
            else:
                return "Document flagged for review - compliance issues detected"
        else:  # uncertain
            return "Document requires manual review - compliance validation uncertain"
    
    async def _assess_content_quality(
        self, 
        file_path: str, 
        document_type: DocumentType
    ) -> ContentQualityScore:
        """
        Step 2: AI-Powered Ingestion & Vetting
        Assess content quality using LLM
        """
        try:
            # Extract text content based on file type
            content = await self._extract_text_content(file_path)
            
            # Define quality assessment prompt
            assessment_prompt = f"""
            Analyze the following agricultural document for quality and scientific accuracy.
            
            Document Type: {document_type.value}
            
            Content:
            {content[:4000]}  # Limit content for token efficiency
            
            Please assess the following criteria (score 0-100):
            1. Scientific Accuracy (25 points)
            2. Factual Language vs Marketing Hype (25 points)
            3. Completeness of Information (25 points)
            4. Clarity and Structure (25 points)
            
            Also identify:
            - Any prohibited marketing language
            - Missing critical information
            - Potential safety concerns
            - Recommendations for improvement
            
            Respond in JSON format:
            {{
                "total_score": <0-100>,
                "scientific_accuracy": <0-25>,
                "factual_language": <0-25>,
                "completeness": <0-25>,
                "clarity": <0-25>,
                "issues": ["issue1", "issue2"],
                "recommendations": ["rec1", "rec2"]
            }}
            """
            
            # Get LLM assessment using optimized service
            from app.services.optimized_llm_service import LLMTask, LLMComplexity
            
            task = LLMTask(
                task_id="quality_assessment",
                prompt=assessment_prompt,
                complexity=LLMComplexity.MEDIUM,
                max_tokens=1000,
                temperature=0.1
            )
            
            result = await self.llm_service.execute_task(task)
            response = result.response
            
            # Parse response
            try:
                assessment_data = json.loads(response)
                return ContentQualityScore(
                    score=assessment_data.get("total_score", 0),
                    issues=assessment_data.get("issues", []),
                    recommendations=assessment_data.get("recommendations", [])
                )
            except json.JSONDecodeError:
                logger.warning("Failed to parse LLM assessment response")
                return ContentQualityScore(score=50.0, issues=["Assessment parsing failed"], recommendations=["Manual review required"])
                
        except Exception as e:
            logger.error(f"Error assessing content quality: {e}")
            return ContentQualityScore(score=0.0, issues=[f"Assessment failed: {str(e)}"], recommendations=["Manual review required"])
    
    async def _extract_text_content(self, file_path: str) -> str:
        """Extract text content from various file formats"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                # Use PyPDF2 or similar for PDF extraction
                import PyPDF2
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()
                return text
                
            elif file_extension in ['.docx', '.doc']:
                # Use python-docx for Word documents
                from docx import Document
                doc = Document(file_path)
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_extension == '.csv':
                # For CSV files, return structure info
                import pandas as pd
                df = pd.read_csv(file_path)
                return f"CSV with {len(df)} rows and columns: {list(df.columns)}"
                
            elif file_extension in ['.txt', '.md']:
                # Plain text files
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
                    
            else:
                return f"Unsupported file type: {file_extension}"
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Error extracting content: {str(e)}"
    
    async def _auto_approve_document(self, document_id: str, db: AsyncSession):
        """Auto-approve high-quality documents"""
        try:
            # Update document status
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            document.processing_status = DocumentStatus.COMPLETED
            document.visibility = "shared"  # Make available to RAG system
            document.organization_metadata["submission_status"] = SubmissionStatus.APPROVED
            document.organization_metadata["approved_at"] = datetime.utcnow().isoformat()
            document.organization_metadata["approved_by"] = "auto_approval"
            
            await db.commit()
            
            # Process document for RAG system
            await self._process_document_for_rag(document, db)
            
            logger.info(f"Document {document_id} auto-approved")
            
        except Exception as e:
            logger.error(f"Error auto-approving document {document_id}: {e}")
            raise
    
    async def _notify_reviewers(self, document_id: str, compliance_result: Dict[str, Any], db: AsyncSession):
        """Notify human reviewers about documents needing review"""
        try:
            # Get super admin users for review
            result = await db.execute(
                select(User).where(User.is_superuser == True)
            )
            reviewers = result.scalars().all()
            
            # Prepare notification data
            decision = compliance_result.get("decision", "UNCERTAIN")
            violations = compliance_result.get("violations", [])
            warnings = compliance_result.get("warnings", [])
            recommendations = compliance_result.get("recommendations", [])
            
            # Create notification message
            if violations:
                message = f"New document submitted requiring review. {len(violations)} regulatory violations found."
            elif warnings:
                message = f"New document submitted requiring review. {len(warnings)} compliance warnings."
            else:
                message = f"New document submitted requiring review. Compliance validation uncertain."
            
            # Send notification to reviewers
            for reviewer in reviewers:
                await self.notification_service.send_notification(
                    user_id=str(reviewer.id),
                    title="Knowledge Base Review Required",
                    message=message,
                    notification_type="knowledge_base_review",
                    metadata={
                        "document_id": document_id,
                        "compliance_decision": decision,
                        "violations": violations,
                        "warnings": warnings,
                        "recommendations": recommendations,
                        "compliance_confidence": compliance_result.get("confidence", 0.0)
                    }
                )
            
            logger.info(f"Review notifications sent for document {document_id}")
            
        except Exception as e:
            logger.error(f"Error notifying reviewers: {e}")
    
    async def approve_document(
        self,
        document_id: str,
        approved_by: str,
        comments: Optional[str] = None,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """
        Step 3: Human Review & Approval
        Approve a document after human review
        """
        try:
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            # Update document status
            document.processing_status = DocumentStatus.COMPLETED
            document.visibility = "shared"
            document.organization_metadata["submission_status"] = SubmissionStatus.APPROVED
            document.organization_metadata["approved_at"] = datetime.utcnow().isoformat()
            document.organization_metadata["approved_by"] = approved_by
            if comments:
                document.organization_metadata["approval_comments"] = comments
            
            await db.commit()
            
            # Process document for RAG system
            await self._process_document_for_rag(document, db)
            
            # Notify submitting organization
            await self._notify_organization_approval(document, db)
            
            return {
                "status": "approved",
                "message": "Document approved and activated in knowledge base"
            }
            
        except Exception as e:
            logger.error(f"Error approving document {document_id}: {e}")
            raise
    
    async def reject_document(
        self,
        document_id: str,
        rejected_by: str,
        reason: str,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """Reject a document with reason"""
        try:
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            # Update document status
            document.processing_status = DocumentStatus.FAILED
            document.organization_metadata["submission_status"] = SubmissionStatus.REJECTED
            document.organization_metadata["rejected_at"] = datetime.utcnow().isoformat()
            document.organization_metadata["rejected_by"] = rejected_by
            document.organization_metadata["rejection_reason"] = reason
            
            await db.commit()
            
            # Notify submitting organization
            await self._notify_organization_rejection(document, reason, db)
            
            return {
                "status": "rejected",
                "message": "Document rejected with reason provided"
            }
            
        except Exception as e:
            logger.error(f"Error rejecting document {document_id}: {e}")
            raise
    
    async def _process_document_for_rag(self, document: KnowledgeBaseDocument, db: AsyncSession):
        """Process approved document for RAG system"""
        try:
            # This would integrate with your existing ChromaDB setup
            # Extract text, create embeddings, and store in vector database
            
            # For now, just mark as processed
            document.chunk_count = 1  # Placeholder
            document.embedding_model = "text-embedding-ada-002"
            
            await db.commit()
            
            logger.info(f"Document {document.id} processed for RAG system")
            
        except Exception as e:
            logger.error(f"Error processing document for RAG: {e}")
            raise
    
    async def approve_document(
        self,
        document_id: str,
        approved_by: str,
        comments: Optional[str] = None,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """
        Superadmin approval of a document
        """
        try:
            logger.info(f"Approving document {document_id} by {approved_by}")
            
            # Get document
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            if document.submission_status != "under_review":
                raise ValueError(f"Document {document_id} is not under review")
            
            # Update document status
            document.submission_status = "approved"
            document.approved_by = approved_by
            document.approved_at = datetime.utcnow()
            document.processing_status = DocumentStatus.COMPLETED
            document.visibility = "shared"  # Make available in RAG
            
            # Update workflow metadata
            document.workflow_metadata.update({
                "approved_at": datetime.utcnow().isoformat(),
                "approved_by": approved_by,
                "approval_comments": comments or "",
                "approval_type": "manual"
            })
            
            await db.commit()
            
            # Process document for RAG system
            await self._process_document_for_rag(document, db)
            
            # Log audit trail
            await self._log_workflow_action(
                document_id=document_id,
                action="approved",
                performed_by=approved_by,
                details={"comments": comments},
                db=db
            )
            
            # Send notification to organization
            await self._send_approval_notification(document, db)
            
            logger.info(f"Document {document_id} approved successfully")
            
            return {
                "success": True,
                "message": "Document approved successfully",
                "document_id": document_id,
                "approved_at": document.approved_at.isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error approving document {document_id}: {e}")
            raise
    
    async def reject_document(
        self,
        document_id: str,
        rejected_by: str,
        reason: str,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """
        Superadmin rejection of a document
        """
        try:
            logger.info(f"Rejecting document {document_id} by {rejected_by}")
            
            # Get document
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            if document.submission_status != "under_review":
                raise ValueError(f"Document {document_id} is not under review")
            
            # Update document status
            document.submission_status = "rejected"
            document.rejection_reason = reason
            document.processing_status = DocumentStatus.FAILED
            
            # Update workflow metadata
            document.workflow_metadata.update({
                "rejected_at": datetime.utcnow().isoformat(),
                "rejected_by": rejected_by,
                "rejection_reason": reason,
                "rejection_type": "manual"
            })
            
            await db.commit()
            
            # Log audit trail
            await self._log_workflow_action(
                document_id=document_id,
                action="rejected",
                performed_by=rejected_by,
                details={"reason": reason},
                db=db
            )
            
            # Send notification to organization
            await self._send_rejection_notification(document, reason, db)
            
            logger.info(f"Document {document_id} rejected successfully")
            
            return {
                "success": True,
                "message": "Document rejected",
                "document_id": document_id,
                "rejection_reason": reason,
                "rejected_at": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error rejecting document {document_id}: {e}")
            raise
    
    async def _log_workflow_action(
        self,
        document_id: str,
        action: str,
        performed_by: str,
        details: Dict[str, Any],
        comments: Optional[str] = None,
        db: AsyncSession = None
    ):
        """Log workflow action to audit trail"""
        try:
            from app.models.knowledge_base import KnowledgeBaseWorkflowAudit
            
            audit_entry = KnowledgeBaseWorkflowAudit(
                document_id=document_id,
                action=action,
                performed_by=performed_by,
                details=details,
                comments=comments
            )
            
            db.add(audit_entry)
            await db.commit()
            
        except Exception as e:
            logger.error(f"Error logging workflow action: {e}")
    
    async def _send_approval_notification(self, document: KnowledgeBaseDocument, db: AsyncSession):
        """Send approval notification to organization"""
        try:
            # Get organization admin users
            from app.models.organization import OrganizationMembership
            result = await db.execute(
                select(OrganizationMembership).where(
                    and_(
                        OrganizationMembership.organization_id == document.organization_id,
                        OrganizationMembership.role.in_(["owner", "admin"]),
                        OrganizationMembership.is_active == True
                    )
                )
            )
            memberships = result.scalars().all()
            
            for membership in memberships:
                await self.notification_service.send_notification(
                    user_id=str(membership.user_id),
                    notification_type="document_approved",
                    title="Document Approved",
                    message=f"Your document '{document.filename}' has been approved and is now available in the knowledge base.",
                    document_id=str(document.id)
                )
                
        except Exception as e:
            logger.error(f"Error sending approval notification: {e}")
    
    async def _send_rejection_notification(self, document: KnowledgeBaseDocument, reason: str, db: AsyncSession):
        """Send rejection notification to organization"""
        try:
            # Get organization admin users
            from app.models.organization import OrganizationMembership
            result = await db.execute(
                select(OrganizationMembership).where(
                    and_(
                        OrganizationMembership.organization_id == document.organization_id,
                        OrganizationMembership.role.in_(["owner", "admin"]),
                        OrganizationMembership.is_active == True
                    )
                )
            )
            memberships = result.scalars().all()
            
            for membership in memberships:
                await self.notification_service.send_notification(
                    user_id=str(membership.user_id),
                    notification_type="document_rejected",
                    title="Document Rejected",
                    message=f"Your document '{document.filename}' was rejected. Reason: {reason}",
                    document_id=str(document.id)
                )
                
        except Exception as e:
            logger.error(f"Error sending rejection notification: {e}")
    
    async def check_expiring_documents(self, db: AsyncSession = None) -> List[Dict[str, Any]]:
        """
        Step 4: Expiration Management
        Check for documents expiring soon and send reminders
        """
        try:
            expiring_docs = []
            
            for days in self.reminder_days:
                expiration_threshold = datetime.utcnow() + timedelta(days=days)
                
                # Find documents expiring in the specified timeframe
                result = await db.execute(
                    select(KnowledgeBaseDocument).where(
                        and_(
                            KnowledgeBaseDocument.processing_status == DocumentStatus.COMPLETED,
                            KnowledgeBaseDocument.visibility == "shared",
                            func.json_extract_path_text(
                                KnowledgeBaseDocument.organization_metadata, 
                                'expiration_date'
                            ).cast(func.date) == expiration_threshold.date()
                        )
                    )
                )
                documents = result.scalars().all()
                
                for doc in documents:
                    # Check if reminder already sent
                    metadata = doc.organization_metadata or {}
                    reminders_sent = metadata.get("reminders_sent", [])
                    
                    if days not in reminders_sent:
                        # Send reminder
                        await self._send_expiration_reminder(doc, days, db)
                        
                        # Mark reminder as sent
                        reminders_sent.append(days)
                        metadata["reminders_sent"] = reminders_sent
                        doc.organization_metadata = metadata
                        
                        expiring_docs.append({
                            "document_id": str(doc.id),
                            "filename": doc.filename,
                            "organization_id": str(doc.organization_id),
                            "days_until_expiration": days,
                            "expiration_date": metadata.get("expiration_date")
                        })
            
            await db.commit()
            return expiring_docs
            
        except Exception as e:
            logger.error(f"Error checking expiring documents: {e}")
            raise
    
    async def _send_expiration_reminder(self, document: KnowledgeBaseDocument, days: int, db: AsyncSession):
        """Send expiration reminder to organization"""
        try:
            # Get organization contact
            result = await db.execute(
                select(Organization).where(Organization.id == document.organization_id)
            )
            organization = result.scalar_one()
            
            # Get organization admin users
            from app.models.organization import OrganizationMembership
            result = await db.execute(
                select(OrganizationMembership).where(
                    and_(
                        OrganizationMembership.organization_id == document.organization_id,
                        OrganizationMembership.role.in_(["owner", "admin"]),
                        OrganizationMembership.is_active == True
                    )
                )
            )
            memberships = result.scalars().all()
            
            for membership in memberships:
                await self.notification_service.send_notification(
                    user_id=str(membership.user_id),
                    title=f"Knowledge Base Content Expiring in {days} days",
                    message=f"Your document '{document.filename}' will expire in {days} days. Please review and renew if needed.",
                    notification_type="knowledge_base_expiration",
                    metadata={
                        "document_id": str(document.id),
                        "days_remaining": days,
                        "expiration_date": document.organization_metadata.get("expiration_date")
                    }
                )
            
            logger.info(f"Expiration reminder sent for document {document.id} ({days} days)")
            
        except Exception as e:
            logger.error(f"Error sending expiration reminder: {e}")
    
    async def renew_document(
        self,
        document_id: str,
        renewed_by: str,
        new_file_path: Optional[str] = None,
        new_expiration_months: Optional[int] = None,
        db: AsyncSession = None
    ) -> Dict[str, Any]:
        """
        Step 5: Renewal Process
        Renew an expiring document
        """
        try:
            result = await db.execute(
                select(KnowledgeBaseDocument).where(KnowledgeBaseDocument.id == document_id)
            )
            document = result.scalar_one()
            
            # Update expiration date
            expiration_months = new_expiration_months or self.default_expiration_months
            new_expiration_date = datetime.utcnow() + timedelta(days=expiration_months * 30)
            
            metadata = document.organization_metadata or {}
            metadata["expiration_date"] = new_expiration_date.isoformat()
            metadata["renewed_at"] = datetime.utcnow().isoformat()
            metadata["renewed_by"] = renewed_by
            metadata["version"] = metadata.get("version", 1) + 1
            metadata["reminders_sent"] = []  # Reset reminders
            
            document.organization_metadata = metadata
            
            # If new file provided, re-process
            if new_file_path:
                document.file_path = new_file_path
                document.processing_status = DocumentStatus.PENDING
                
                # Re-assess quality
                quality_assessment = await self._assess_content_quality(new_file_path, document.document_type)
                metadata["quality_score"] = quality_assessment.score
                metadata["quality_issues"] = quality_assessment.issues
                metadata["quality_recommendations"] = quality_assessment.recommendations
                
                if quality_assessment.is_acceptable(self.quality_threshold):
                    document.processing_status = DocumentStatus.COMPLETED
                    metadata["submission_status"] = SubmissionStatus.APPROVED
                else:
                    metadata["submission_status"] = SubmissionStatus.UNDER_REVIEW
            
            await db.commit()
            
            return {
                "status": "renewed",
                "new_expiration_date": new_expiration_date.isoformat(),
                "version": metadata["version"],
                "message": "Document renewed successfully"
            }
            
        except Exception as e:
            logger.error(f"Error renewing document {document_id}: {e}")
            raise
    
    async def deactivate_expired_documents(self, db: AsyncSession = None) -> List[str]:
        """Deactivate expired documents from RAG system"""
        try:
            today = datetime.utcnow().date()
            
            result = await db.execute(
                select(KnowledgeBaseDocument).where(
                    and_(
                        KnowledgeBaseDocument.processing_status == DocumentStatus.COMPLETED,
                        KnowledgeBaseDocument.visibility == "shared",
                        func.json_extract_path_text(
                            KnowledgeBaseDocument.organization_metadata, 
                            'expiration_date'
                        ).cast(func.date) < today
                    )
                )
            )
            expired_docs = result.scalars().all()
            
            deactivated_ids = []
            for doc in expired_docs:
                # Deactivate from RAG system
                doc.visibility = "internal"  # Remove from shared knowledge
                doc.organization_metadata["submission_status"] = SubmissionStatus.EXPIRED
                doc.organization_metadata["deactivated_at"] = datetime.utcnow().isoformat()
                
                deactivated_ids.append(str(doc.id))
            
            await db.commit()
            
            logger.info(f"Deactivated {len(deactivated_ids)} expired documents")
            return deactivated_ids
            
        except Exception as e:
            logger.error(f"Error deactivating expired documents: {e}")
            raise
    
    async def _notify_organization_approval(self, document: KnowledgeBaseDocument, db: AsyncSession):
        """Notify organization of document approval"""
        # Implementation for approval notification
        pass
    
    async def _notify_organization_rejection(self, document: KnowledgeBaseDocument, reason: str, db: AsyncSession):
        """Notify organization of document rejection"""
        # Implementation for rejection notification
        pass
